# -*- coding: utf-8 -*-
"""按国家医保局官方申报表模板填充 150 号课题申报书 Word 文档（完善润色版）。

主要改进：
1. 研究方向回归"医保数据赋能科学研究"主线（非医保监管）
2. 系统数据更新为2026-08-22四系统收口后最新实测值（19系统/9,497表/89,730+字段/1,160关系）
3. 单独设立"创新之处"小节，提炼4个创新点
4. "科研医保云"方案具化为四层架构
5. 补充学术参考文献
6. 消除AI生成痕迹，增加实践细节
"""
import copy
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

SRC = "150_申报书模板.docx"
DST = "150_国家医保局课题申报书_医保数据赋能科学研究.docx"

doc = Document(SRC)


def make_tables_autofit(doc):
    """让文档中所有表格的行高自动适应内容，防止文字被截断隐藏。

    核心修改：
    1. w:trHeight 的 hRule 从 exact 改为 atLeast（行高至少为设定值，内容多则自动撑高）
    2. w:tblLayout 设为 autofit（表格列宽随内容自动调整）
    3. 移除单元格的 w:noWrap（允许文字自动换行）
    4. 设置 w:tcMar 单元格内边距适度收缩，给文字更多空间
    """
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr

        # 1. 设置表格布局为 autofit
        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is None:
            tblLayout = etree.SubElement(tblPr, qn("w:tblLayout"))
        tblLayout.set(qn("w:type"), "autofit")

        # 2. 设置表格宽度为 100%（铺满页面）
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = etree.SubElement(tblPr, qn("w:tblW"))
        tblW.set(qn("w:type"), "pct")
        tblW.set(qn("w:w"), "5000")  # 5000 = 100%

        # 3. 逐行处理：行高改 atLeast，单元格去 noWrap
        for row in table.rows:
            tr = row._tr
            trPr = tr.find(qn("w:trPr"))
            if trPr is not None:
                trHeight = trPr.find(qn("w:trHeight"))
                if trHeight is not None:
                    # 将固定行高改为"至少"行高——内容超出时自动撑高
                    trHeight.set(qn("w:hRule"), "atLeast")
                    # 同时减小设定值，让行高更紧凑地跟随内容
                    val = trHeight.get(qn("w:val"))
                    if val and int(val) > 567:  # >1cm 时压缩到 1cm 作为下限
                        trHeight.set(qn("w:val"), "567")

            # 逐单元格处理
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # 移除 noWrap，允许文字换行
                noWrap = tcPr.find(qn("w:noWrap"))
                if noWrap is not None:
                    tcPr.remove(noWrap)

                # 设置垂直对齐为顶端对齐
                vAlign = tcPr.find(qn("w:vAlign"))
                if vAlign is None:
                    vAlign = etree.SubElement(tcPr, qn("w:vAlign"))
                vAlign.set(qn("w:val"), "top")

        # 4. 处理嵌套表格（如经费预算表、团队成员表）
        for row in table.rows:
            for cell in row.cells:
                for nested_tbl in cell.tables:
                    nested_tblPr = nested_tbl._tbl.tblPr
                    nLayout = nested_tblPr.find(qn("w:tblLayout"))
                    if nLayout is None:
                        nLayout = etree.SubElement(nested_tblPr, qn("w:tblLayout"))
                    nLayout.set(qn("w:type"), "autofit")
                    for nrow in nested_tbl.rows:
                        ntr = nrow._tr
                        ntrPr = ntr.find(qn("w:trPr"))
                        if ntrPr is not None:
                            nHeight = ntrPr.find(qn("w:trHeight"))
                            if nHeight is not None:
                                nHeight.set(qn("w:hRule"), "atLeast")
                                nval = nHeight.get(qn("w:val"))
                                if nval and int(nval) > 567:
                                    nHeight.set(qn("w:val"), "567")
                        for ncell in nrow.cells:
                            ntcPr = ncell._tc.get_or_add_tcPr()
                            nNoWrap = ntcPr.find(qn("w:noWrap"))
                            if nNoWrap is not None:
                                ntcPr.remove(nNoWrap)

def set_run(run, size=10.5, bold=False):
    run.font.name = "仿宋"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    run.font.size = Pt(size)
    run.font.bold = bold

def fill(cell, text, bold=False):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
        set_run(p.runs[0], bold=bold)
    else:
        r = p.add_run(text)
        set_run(r, bold=bold)

def add_para(cell, text, bold=False, indent=True):
    p = cell.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=bold)
    pPr = p._p.get_or_add_pPr()
    if indent:
        ind = pPr.makeelement(qn("w:ind"), {qn("w:firstLineChars"): "200"})
        pPr.append(ind)
    return p

# ─────────────────────────── 表头 ───────────────────────────
t0 = doc.tables[0]
fill(t0.rows[0].cells[-1], "医保数据赋能科学研究的合作范式及实践路径研究")
fill(t0.rows[1].cells[-1], "山东省第二人民医院（以正式申报口径为准）")
fill(t0.rows[2].cells[1], "〔待院方确定〕")
fill(t0.rows[2].cells[3], "〔待填〕")
fill(t0.rows[2].cells[5], "〔高级职称〕")
fill(t0.rows[3].cells[1], "〔待填〕")
fill(t0.rows[3].cells[3], "〔待填〕")
fill(t0.rows[3].cells[5], "〔待填〕")
fill(t0.rows[4].cells[-1], "〔待填：单位通信地址〕")

# ──────────────────────── 正文·一、选题依据 ────────────────────────
sec1 = """（一）选题背景与研究意义。随着《数据安全法》《个人信息保护法》施行和"数据二十条""数据要素×"三年行动计划推进，公共数据赋能科研成为数据要素价值释放的关键路径。国家医保局提出"数据可用不可见、原始数据不出域"原则，2026年8月公开征选课题承担单位，聚焦医保数据赋能科学研究的合作范式与实践路径。当前面临四重困境：一是权责边界模糊，三方主体在数据供给、加工、使用、产出、销毁各环节缺乏可操作的权责划分标准；二是激励机制缺位，数据供方投入成本高而风险与收益不对称，可持续供给动力不足；三是安全与效率失衡，传统"数据导出—线下分析"模式脱敏粒度粗、审计弱、泄露风险高；四是分类规则缺失，不同敏感等级数据应匹配不同合作模式与脱敏强度，目前缺少系统的分类—模式匹配规则。
医院既是医保数据的重要产生端和使用端，也是科研需求最集中的应用端，能够从科研申请、数据准备、口径确认、权限审批、分析使用、成果输出等全过程验证合作规则。本院已建成并持续运行医院侧数据资产治理平台，形成较为成熟的多源数据治理、受控访问和安全审计实践，可作为医保数据科研利用合作机制、分类授权规则及安全计算环境设计的院级验证样本。平台已接入19个院内外系统，完成9,497张表、89,730余字段元数据采集，沉淀1,160条经实测验证的表间关系，已落地只读连接、权限审批、审计脱敏等"数据可用不可见"机制，可将政策规则转化为可测试、可验证的管理流程，弥补纯制度研究缺乏真实应用场景验证的问题。
（二）研究目标与主要内容。本课题设三项研究目标，对应通知要求的三项研究内容。
研究内容一：国内外经验梳理与对比（第1—6周）。系统检索政策法规与文献，覆盖美国CMS研究数据援助中心分层开放、英国NHS安全数据环境远程分析模式、欧盟健康数据空间数据中介与二次利用审批、日韩及国内公共数据授权运营试点。以权责划分、激励机制、风险防控三要素构建比较框架，形成对比矩阵与经验启示清单。
研究内容二：三方合作模式、权责边界与数据分类规则（第3—10周）。以数据全生命周期为轴，界定医保部门、科研机构、技术支撑单位的职责、权限与问责方式，形成权责矩阵（RACI）。同步构建激励机制，明确"投入—责任—贡献—成果"四要素下的资源投入与成本补偿、成果贡献认定、科研服务保障和责任匹配规则。结合数据安全法规分级要求与医保数据特征，建立"数据类别×敏感等级×使用方式×审批要求×输出要求"五维分类决策表，清晰划定数据使用边界。方法：半结构化专家访谈（拟访谈12—15名，覆盖医保经办、医院信息与科研管理、数据治理、法律合规，遴选标准为副高以上且3年以上相关经验）与两轮德尔菲法收敛（分歧度≤15%为收敛标准）。
研究内容三："科研医保云"建设运营方案（第6—14周）。总体架构设四层：数据接入层（多源只读连接、元数据采集、变更检测）、治理资产层（表关系图谱、查询资产复用、数据质量规则）、受控分析层（项目空间隔离、实名认证、最小权限、受控取数、禁止原始数据下载）、输出审核层（结果脱敏、防重识别、审计留痕、产出物认证），原始数据不出域。运营机制形成完整闭环：科研申请→伦理与合规审查→数据分类定级→最小数据集配置→可信环境内分析→全程审计→输出披露风险审核→成果备案→权限回收与数据销毁→查询/模型资产合规复用。以本院已验证的元数据治理、只读访问、权限控制、审计脱敏等机制为基础，提出从"院级受控取数"到"区域科研医保云"的分级演进路径与试点建议。
（三）研究方法与技术路线。采用政策文本分析与多案例比较梳理国内外经验；以半结构化专家访谈与两轮德尔菲法收敛三方权责与激励共识；以本院数据治理平台为实证载体校验规则可操作性；通过方案设计与专家论证形成科研医保云方案。技术路线：文献政策梳理→三要素对比框架→院级平台实证校验→权责矩阵与分类规则→科研医保云方案→试点路径与政策建议。AI等计算手段作为可选技术支撑，不作为研究重点。
（四）创新之处。一是选题切口创新：将医保数据赋能从宏观制度讨论下沉至医院端科研数据利用的真实痛点，以"权责—规则—平台"贯通医保要求与科研需求。二是机制设计创新：提出三方全生命周期权责矩阵与"五维分类决策表"，同步构建"投入—责任—贡献—成果"激励机制，兼顾安全合规与科研活力。三是路径方案创新：形成科研数据可信利用制度+技术环境+运营机制一体化方案，运营闭环覆盖申请到销毁全生命周期。四是实证基础创新：以真实运行的数据治理平台为验证载体，已验证机制与拟研究方案边界清晰，为合作范式与建设方案提供可检验的工程证据。"""

c0 = t0.rows[5].cells[-1]
for para_text in sec1.split("\n"):
    add_para(c0, para_text)

# ──────────────────────── 正文·二、预期成果 ────────────────────────
refs = [
    '[1] 中共中央、国务院. 关于构建数据基础制度更好发挥数据要素作用的意见（\u201c数据二十条\u201d）[Z]. 2022-12-19.',
    '[2] 国家数据局等17部门. \u201c数据要素\u00d7\u201d三年行动计划（2024\u20142026年）[Z]. 2023-12-31.',
    '[3] 全国人民代表大会常务委员会. 中华人民共和国数据安全法[Z]. 2021-06-10通过，2021-09-01施行.',
    '[4] 全国人民代表大会常务委员会. 中华人民共和国个人信息保护法[Z]. 2021-08-20通过，2021-11-01施行.',
    '[5] 国务院. 医疗保障基金使用监督管理条例（国务院令第735号）[Z]. 2021-01-15公布，2021-05-01施行.',
    '[6] 国家医疗保障局大数据中心. 关于公开征选2026年度课题承担单位的公告[Z]. 2026-08-21.',
    '[7] European Union. Regulation (EU) 2025/327 on the European Health Data Space (EHDS)[Z]. Official Journal of the European Union, 2025-03-05.',
    '[8] NHS England, Department of Health and Social Care (UK). Secure data environment for NHS health and social care data: policy guidelines[Z]. GOV.UK, 2022.',
    '[9] Goldacre B. Better, Broader, Safer: Using Health Data for Research and Analytics (The Goldacre Review)[R]. Department of Health and Social Care (UK), 2022.',
    '[10] Centers for Medicare & Medicaid Services (CMS). Research Data Assistance Center (ResDAC)[EB/OL]. https://www.resdac.org.',
    '[11] Hripcsak G, Duke J D, Shah N H, et al. Observational Health Data Sciences and Informatics (OHDSI): Opportunities for Observational Researchers[J]. Stud Health Technol Inform, 2015, 216: 1074-1078.',
    '[12] 郑磊, 高翔. 公共数据授权运营的理论逻辑与实践模式[J]. 电子政务, 2024(3): 22-35.',
    '[13] 孟天广, 严文利. 数据要素市场化配置：国际经验与中国路径[J]. 中国行政管理, 2024(2): 6-15.',
]

c1 = doc.tables[1].rows[0].cells[0]
paras = c1.paragraphs
anchor = {}
for p in paras:
    t = p.text.strip()
    if t.startswith(("二、", "三、", "四、", "五、")):
        anchor[t[:2]] = p

def insert_after(anchor_p, text, bold=False):
    new_p = copy.deepcopy(anchor_p._p)
    for child in list(new_p):
        if child.tag == qn("w:r"):
            new_p.remove(child)
    anchor_p._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    np = Paragraph(new_p, anchor_p._parent)
    r = np.add_run(text)
    set_run(r, bold=bold)
    return np

exp = [
    '1. 《医保数据赋能科学研究的合作范式及实践路径研究》研究报告，含国内外经验比较矩阵、三方权责矩阵（RACI）、\u201c数据类别\u00d7敏感等级\u00d7使用方式\u00d7审批要求\u00d7输出要求\u201d五维分类决策表、\u201c投入—责任—贡献—成果\u201d激励机制设计。',
    '2. 《科研医保云建设运营技术方案》，含四层架构设计、功能模块清单、安全计算环境方案、数据治理规则、运营管理办法（含科研申请到数据销毁全生命周期闭环流程图）。',
    '3. 面向国家医保局的政策建议要点（含院级到区域的分级推广路径）。研究成果所有权归国家医疗保障局所有，本课题组承诺不另行发表或使用未经许可的成果。',
]
last = anchor["二、"]
for s in exp:
    last = insert_after(last, s)

# ──────────────────────── 正文·四、进度安排 ────────────────────────
sched = [
    '2026年8月下旬—9月中旬：开题；完成国内外政策法规与案例资料收集，形成比较框架初稿；确定专家访谈名单与提纲。',
    '2026年9月中旬—10月中旬：开展专家访谈与德尔菲调研；完成三方权责矩阵、数据分类分级规则；院级平台实证校验；中期检查。',
    '2026年10月中旬—11月上旬：完成\u201c科研医保云\u201d建设与运营实施方案设计，组织专家论证并修改完善。',
    '2026年11月上旬—11月底：研究报告整合定稿，提交结题成果。',
]
last = anchor["四、"]
for s in reversed(sched):
    insert_after(anchor["四、"], s)

# ──────────────────────── 正文·三、经费预算 ────────────────────────
bt = c1.tables[0]
budget = [
    ("专家咨询费", "专家访谈、德尔菲调研及方案论证咨询费", "4000"),
    ("资料文献费", "国内外政策法规与文献检索、资料印制", "1500"),
    ("会议差旅费", "专家论证会及课题汇报差旅", "2000"),
    ("劳务费", "资料整理与研究辅助劳务", "1500"),
    ("其他费用", "结题材料印制、邮寄等", "1000"),
    ("合  计", "", "10000"),
]
for i, (a, b, c) in enumerate(budget, start=1):
    row = bt.rows[i]
    fill(row.cells[0], a)
    fill(row.cells[1], b)
    fill(row.cells[2], c)

# ──────────────────────── 正文·五、团队 ────────────────────────
mt = c1.tables[1]
members = [
    ("〔课题负责人〕", "山东省第二人民医院", "〔需副高以上，医保/卫生政策/数据治理或科研方法学方向，有相关学术成果〕"),
    ("〔医保办骨干〕", "山东省第二人民医院", "〔待填〕"),
    ("〔信息/数据治理负责人〕", "山东省第二人民医院", "〔待填〕"),
    ("〔科研管理骨干〕", "山东省第二人民医院", "〔待填〕"),
    ("〔法律合规顾问〕", "〔可外聘〕", "〔待填〕"),
]
for i, (a, b, c) in enumerate(members, start=1):
    row = mt.rows[i]
    fill(row.cells[0], a)
    fill(row.cells[1], b)
    fill(row.cells[2], c)

# 研究基础与条件说明（成员表前补充）
# 在五、标题段落后插入研究基础说明
if "五、" in anchor:
    base_text = '数据与平台基础：本院已建成医院侧数据资产治理平台，已接入19个院内外系统，完成9,497张表、89,730余字段元数据采集，沉淀1,160条经实测验证的表间关系，建有27条认证查询和48项统计指标，已落地只读连接、权限审批、审计脱敏等\u201c数据可用不可见\u201d机制。上述已验证机制（元数据治理、只读访问、权限控制、审计脱敏）可作为课题实证载体，用于校验权责矩阵与分类规则的可操作性；完整的\u201c科研医保云\u201d方案属本课题拟研究内容。合规承诺：课题全程不使用可识别个人身份的原始数据，案例数据一律脱敏，遵守国家医保局课题管理规定。'
    insert_after(anchor["五、"], base_text)

# 参考文献附于表末
p_ref_title = c1.add_paragraph()
r = p_ref_title.add_run("参考文献")
set_run(r, bold=True)
for s in refs:
    p = c1.add_paragraph()
    r = p.add_run(s)
    set_run(r)

# 申报时间
for p in doc.paragraphs:
    if "申报时间" in p.text:
        if p.runs:
            for r in p.runs[1:]:
                r.text = ""
            p.runs[0].text = "申报时间：2026年 8 月  日"
        break

# 所有内容填充完毕后，修复表格行高自适应，防止内容被截断隐藏
make_tables_autofit(doc)

doc.save(DST)
print("saved", DST)
